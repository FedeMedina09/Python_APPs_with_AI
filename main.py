import tkinter as tk
from tkinter import filedialog
import docx
import txt_Prompt

from groq import Groq

# Constantes de color
AZUL = "#3498db"
BLANCO = "#ffffff"
GRIS = "#f0f0f0"

# Funciones
def cancelar():
    ventana.destroy()

def seleccionar_ruta():
    ruta_seleccionada = filedialog.askdirectory()
    ruta.set(ruta_seleccionada)

# Función para mostrar la descripción
def mostrar_descripcion(descripcion):
    descripcion_ventana = tk.Toplevel(ventana)
    descripcion_ventana.title("Descripción")
    descripcion_label = tk.Label(descripcion_ventana, text=descripcion, font=("Arial", 12), wraplength=400)
    descripcion_label.pack(padx=20, pady=20)

def crear():
    try:
        client = Groq(
            api_key=("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"),
        )

        # Obtener los valores de los inputs
        input1 = entrada1.get()
        input2 = entrada2.get()
        input3 = entrada3.get()
        input4 = entrada4.get()
        input5 = entrada5.get()
        input6 = entrada6.get()
        input7 = entrada7.get()
        input8 = entrada8.get()
        input9 = entrada9.get()
        input10 = entrada10.get()
        ruta_archivo = ruta.get()

        # Crear el contenido del prompt
        contenido = txt_Prompt.prompt + """
            Marco de negocio: {}
            Alcance de la Robotización: {}
            Sistema de Destino: {}
            Área Impactada: {}
            Precondición/Archivo de Entrada: {}
            Detalle de la Solución de Robotización: {}
            Políticas de casos duplicados: {}
            Reporte: {}
            Excepciones de negocio: {}
            Transcripcion Completa: {}
        """.format(input1, input2, input3, input4, input5, input6, input7, input8, input9, input10)

        # Guardar ruta de destino
        ruta_destino = "{}".format(ruta_archivo)

        # Crear el chat completion
        chat_completion = client.chat.completions.create(
            messages=[
                {
                    "role": "user",
                    "content": contenido,
                }
            ],
            model="llama-3.2-90b-text-preview",
        )
        temperature=0.5,
        max_tokens=8192,
        top_p=1,
        stream=True,
        stop=None,

        # Imprimir el resultado
        #print(chat_completion.choices[0].message.content)
        ventana.destroy()
        #print(ruta_destino)

        # Guardar la respuesta en un archivo de Word
        documento = docx.Document()
        documento.add_heading("PDD Preliminar", 0)
        documento.add_paragraph(chat_completion.choices[0].message.content)
        documento.save(ruta_destino+"/PDD Preliminar.docx")

        # Imprimir el resultado
        #print("La respuesta se ha guardado en un archivo de Word llamado 'PDD.docx'")

    except Exception as e:
        print(f"Error: {e}")

# Crear la ventana
ventana = tk.Tk()
ventana.title("RPA PDD Maker")
ventana.geometry("1000x600")
ventana.configure(bg=GRIS)

# Crear un marco alrededor de la ventana
marco_ventana = tk.Frame(ventana, bg=AZUL, highlightbackground=AZUL, highlightthickness=2)
marco_ventana.grid(row=0, column=0, sticky="nsew")

# Crear un marco alrededor de la ventana
marco_ventana = tk.Frame(ventana, bg=AZUL, highlightbackground=AZUL, highlightthickness=2)
marco_ventana.grid(row=0, column=1, sticky="nsew")

# Crear un marco alrededor de la ventana
marco_ventana = tk.Frame(ventana, bg=AZUL, highlightbackground=AZUL, highlightthickness=2)
marco_ventana.grid(row=0, column=2, sticky="nsew")

# Crear un marco alrededor de la ventana
marco_ventana = tk.Frame(ventana, bg=AZUL, highlightbackground=AZUL, highlightthickness=2, width=600)
marco_ventana.grid(row=0, column=3, sticky="nsew")

# Crear el título
titulo = tk.Label(ventana, text="               Información requerida para PDD", font=("Arial", 24), bg=AZUL, fg=BLANCO)
titulo.grid(row=0, column=0, columnspan=3, padx=0, pady=0)

# Crear boton de ayuda
info_button = tk.Button(ventana, text="?", font=("Arial", 8), bg=AZUL, fg=BLANCO, command=lambda: mostrar_descripcion("""
                                                                                                                       Marco de Negocio: Breve resumen de lo que realizan hoy los usuarios, sistemas y áreas que intervienen. Es importante que quede claro desde dónde se parte y qué se obtiene y para qué (el valor agregado del proceso de negocio).

                                                                                                                       Alcance de la Robotización: Qué parte del proceso se robotiza (desde dónde hasta dónde), y qué queda a cargo del usuario.

                                                                                                                       Área Impactada: Dirección Médica a la que impacta la Robotización.

                                                                                                                       Precondición/Archivo de Entrada: Detalle de los datos de cada una de las columnas (nombre, tipo, obligatoriedad, valor default si es opcional, valores permitidos).

                                                                                                                       Detalle de la Solución de Robotización del Proceso de Negocio: Paso a paso de lo que hace el bot. Interfaces que utiliza, validaciones que realiza, decisiones que toma, curso de navegación de la aplicación, caso exitoso y casos no exitosos.

                                                                                                                       Excepciones de Negocio: Listado de Excepciones de Negocio y breve descripción si es necesario.
                                                                                                                       """))
info_button.grid(row=0, column=2, padx=0, pady=0)

# Crear los inputs
entrada1 = tk.Entry(ventana, font=("Arial", 12), width=50, bg=BLANCO, fg=AZUL)
entrada1.grid(row=1, column=1, padx=10, pady=10)
tk.Label(ventana, text="Marco de Negocio", font=("Arial", 12), bg=GRIS).grid(row=1, column=0, padx=20, pady=10)

entrada2 = tk.Entry(ventana, font=("Arial", 12), width=50, bg=BLANCO, fg=AZUL)
entrada2.grid(row=2, column=1, padx=10, pady=10)
tk.Label(ventana, text="Alcance de la Robotización", font=("Arial", 12), bg=GRIS).grid(row=2, column=0, padx=20, pady=10)

entrada3 = tk.Entry(ventana, font=("Arial", 12), width=50, bg=BLANCO, fg=AZUL)
entrada3.grid(row=3, column=1, padx=10, pady=10)
tk.Label(ventana, text="Sistema de Destino", font=("Arial", 12), bg=GRIS).grid(row=3, column=0, padx=20, pady=10)

entrada4 = tk.Entry(ventana, font=("Arial", 12), width=50, bg=BLANCO, fg=AZUL)
entrada4.grid(row=4, column=1, padx=10, pady=10)
tk.Label(ventana, text="Área Impactada", font=("Arial", 12), bg=GRIS).grid(row=4, column=0, padx=20, pady=10)

entrada5 = tk.Entry(ventana, font=("Arial", 12), width=50, bg=BLANCO, fg=AZUL)
entrada5.grid(row=5, column=1, padx=10, pady=10)
tk.Label(ventana, text="Precondición/Archivo de Entrada", font=("Arial", 12), bg=GRIS).grid(row=5, column=0, padx=20, pady=10)

entrada6 = tk.Entry(ventana, font=("Arial", 12), width=50, bg=BLANCO, fg=AZUL)
entrada6.grid(row=6, column=1, padx=10, pady=10)
tk.Label(ventana, text="Detalle de la Solución de Robotización", font=("Arial", 12), bg=GRIS).grid(row=6, column=0, padx=20, pady=10)

entrada7 = tk.Entry(ventana, font=("Arial", 12), width=50, bg=BLANCO, fg=AZUL)
entrada7.grid(row=7, column=1, padx=10, pady=10)
tk.Label(ventana, text="Políticas de casos duplicados", font=("Arial", 12), bg=GRIS).grid(row=7, column=0, padx=20, pady=10)

entrada8 = tk.Entry(ventana, font=("Arial", 12), width=50, bg=BLANCO, fg=AZUL)
entrada8.grid(row=8, column=1, padx=10, pady=10)
tk.Label(ventana, text="Reporte", font=("Arial", 12), bg=GRIS).grid(row=8, column=0, padx=20, pady=10)

entrada9 = tk.Entry(ventana, font=("Arial", 12), width=50, bg=BLANCO, fg=AZUL)
entrada9.grid(row=9, column=1, padx=10, pady=10)
tk.Label(ventana, text="Excepciones de negocio", font=("Arial", 12), bg=GRIS).grid(row=9, column=0, padx=20, pady=10)

entrada10 = tk.Entry(ventana, font=("Arial", 12), width=50, bg=BLANCO, fg=AZUL)
entrada10.grid(row=10, column=1, padx=10, pady=10)
tk.Label(ventana, text="Transcripcion Completa", font=("Arial", 12), bg=GRIS).grid(row=10, column=0, padx=20, pady=10)

# Crear el input para la ruta del archivo
ruta = tk.StringVar()
ruta_label = tk.Label(ventana, text="Ruta de Destino de PDD", font=("Arial", 12), bg=GRIS)
ruta_label.grid(row=11, column=0, padx=20, pady=10)
ruta_entry = tk.Entry(ventana, textvariable=ruta, font=("Arial", 12), width=50, bg=BLANCO, fg=AZUL)
ruta_entry.grid(row=11, column=1, padx=20, pady=10)
ruta_button = tk.Button(ventana, text="Seleccionar ruta", font=("Arial", 12), command=seleccionar_ruta, bg=AZUL, fg=BLANCO)
ruta_button.grid(row=11, column=2, padx=20, pady=10)

# Crear el botón de cancelar
cancelar_button = tk.Button(ventana, text="Cancelar", font=("Arial", 12), command=cancelar, bg=AZUL, fg=BLANCO)
cancelar_button.grid(row=12, column=0, padx=20, pady=10)

# Crear el botón de crear
crear_button = tk.Button(ventana, text="Generar PDD", font=("Arial", 12), command=crear, bg=AZUL, fg=BLANCO)
crear_button.grid(row=12, column=1, padx=20, pady=10)

# Crear el texto Created By
created_by = tk.Label(ventana, text="Created by Federico Medina", font=("Arial", 10), bg=AZUL, fg=BLANCO)
created_by.grid(row=0, column=0, columnspan=1, padx=0, pady=0)

# Se Agregan comentarios para probar Branch del Repositorio de GitHub
# Iniciar la ventana
ventana.mainloop()